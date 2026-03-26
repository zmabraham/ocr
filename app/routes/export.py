"""Export corrected text routes"""
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from typing import Optional
import io
import csv
import json
from datetime import datetime

from ..database import get_db
from ..models import Document, ErrorRecord

router = APIRouter()


@router.get("/document/{document_id}/text", response_model=dict)
async def export_text(document_id: str, db: Session = Depends(get_db)):
    """Export corrected text as plain text"""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    # Use final_text if available, otherwise ocr_text
    text = document.final_text or document.ocr_text or ""

    return {
        "document_id": document_id,
        "filename": document.filename,
        "export_date": datetime.utcnow().isoformat(),
        "text": text
    }


@router.get("/document/{document_id}/download")
async def download_text(document_id: str, db: Session = Depends(get_db)):
    """Download corrected text as a .txt file"""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    text = document.final_text or document.ocr_text or ""
    filename = f"{document.filename}_corrected.txt"

    # Create file in memory
    buffer = io.BytesIO()
    buffer.write(text.encode('utf-8'))
    buffer.seek(0)

    return StreamingResponse(
        buffer,
        media_type="text/plain",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/document/{document_id}/corrections-log")
async def export_corrections_log(document_id: str, db: Session = Depends(get_db)):
    """Export a log of all corrections made to the document"""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    errors = db.query(ErrorRecord).filter(
        ErrorRecord.document_id == document_id,
        ErrorRecord.status.in_(["corrected", "approved", "skipped"])
    ).order_by(ErrorRecord.position).all()

    log_entries = []
    for error in errors:
        entry = {
            "position": error.position,
            "original_word": error.original_word,
            "confidence": error.confidence,
            "context": error.context,
            "status": error.status,
            "reviewed_at": error.reviewed_at.isoformat() if error.reviewed_at else None,
            "reviewed_by": error.reviewed_by
        }

        if error.status == "corrected":
            if error.custom_correction:
                entry["correction"] = error.custom_correction
                entry["correction_type"] = "custom"
            elif error.selected_correction is not None and error.suggestions:
                entry["correction"] = error.suggestions[error.selected_correction]
                entry["correction_type"] = "suggested"
        elif error.status == "approved":
            entry["correction"] = error.original_word
            entry["correction_type"] = "approved_as_is"

        log_entries.append(entry)

    return {
        "document_id": document_id,
        "filename": document.filename,
        "export_date": datetime.utcnow().isoformat(),
        "total_corrections": len(log_entries),
        "corrections": log_entries
    }


@router.get("/document/{document_id}/corrections-csv")
async def download_corrections_csv(document_id: str, db: Session = Depends(get_db)):
    """Download corrections log as CSV file"""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    errors = db.query(ErrorRecord).filter(
        ErrorRecord.document_id == document_id,
        ErrorRecord.status.in_(["corrected", "approved", "skipped"])
    ).order_by(ErrorRecord.position).all()

    # Create CSV in memory
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow([
        "Position", "Original Word", "Confidence", "Context",
        "Status", "Correction", "Correction Type", "Reviewed At", "Reviewed By"
    ])

    for error in errors:
        correction = None
        correction_type = None

        if error.status == "corrected":
            if error.custom_correction:
                correction = error.custom_correction
                correction_type = "custom"
            elif error.selected_correction is not None and error.suggestions:
                correction = error.suggestions[error.selected_correction]
                correction_type = "suggested"
        elif error.status == "approved":
            correction = error.original_word
            correction_type = "approved_as_is"

        writer.writerow([
            error.position,
            error.original_word,
            error.confidence,
            error.context,
            error.status,
            correction,
            correction_type,
            error.reviewed_at.isoformat() if error.reviewed_at else "",
            error.reviewed_by or ""
        ])

    # Convert to bytes
    buffer.seek(0)
    csv_bytes = buffer.getvalue().encode('utf-8')
    byte_buffer = io.BytesIO(csv_bytes)
    byte_buffer.seek(0)

    filename = f"{document.filename}_corrections.csv"

    return StreamingResponse(
        byte_buffer,
        media_type="text/csv",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/document/{document_id}/statistics")
async def get_document_statistics(document_id: str, db: Session = Depends(get_db)):
    """Get detailed statistics about the document and corrections"""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    errors = db.query(ErrorRecord).filter(ErrorRecord.document_id == document_id).all()

    total = len(errors)
    corrected = sum(1 for e in errors if e.status == "corrected")
    approved = sum(1 for e in errors if e.status == "approved")
    skipped = sum(1 for e in errors if e.status == "skipped")
    pending = sum(1 for e in errors if e.status == "pending")

    # Calculate confidence distribution
    if errors:
        avg_confidence = sum(e.confidence for e in errors) / total
        low_confidence = sum(1 for e in errors if e.confidence < 0.70)
        medium_confidence = sum(1 for e in errors if 0.70 <= e.confidence < 0.85)
        high_confidence = sum(1 for e in errors if e.confidence >= 0.85)
    else:
        avg_confidence = 0
        low_confidence = medium_confidence = high_confidence = 0

    return {
        "document_id": document_id,
        "filename": document.filename,
        "overall_confidence": document.overall_confidence,
        "hebrew_percentage": document.hebrew_percentage,
        "total_pages": document.total_pages,
        "error_statistics": {
            "total_errors": total,
            "corrected": corrected,
            "approved": approved,
            "skipped": skipped,
            "pending": pending,
            "completion_rate": round(((total - pending) / total * 100) if total > 0 else 100, 2)
        },
        "confidence_statistics": {
            "average_confidence": round(avg_confidence, 3),
            "low_confidence_count": low_confidence,
            "medium_confidence_count": medium_confidence,
            "high_confidence_count": high_confidence
        }
    }


@router.get("/document/{document_id}/word")
async def export_word(document_id: str, db: Session = Depends(get_db)):
    """Export corrected text as formatted Word document with Hebrew RTL support"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from bidi.algorithm import get_display
    except ImportError:
        raise HTTPException(
            status_code=503,
            detail="Word export not available. Please install python-docx and python-bidi."
        )

    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    # Get all error records for this document
    errors = db.query(ErrorRecord).filter(ErrorRecord.document_id == document_id).all()

    # Create a map of position -> correction for quick lookup
    corrections_map = {}
    for error in errors:
        if error.status == "corrected":
            if error.custom_correction:
                corrections_map[error.position] = error.custom_correction
            elif error.selected_correction is not None and error.suggestions:
                corrections_map[error.position] = error.suggestions[error.selected_correction]

    # Create Word document
    doc = Document()

    # Set up default style for Hebrew (RTL)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # Add title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_run = title.add_run(document.filename)
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Add metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta.add_run(f"Processed: {document.upload_date.strftime('%Y-%m-%d %H:%M')}\n")
    meta.add_run(f"Pages: {document.total_pages or 'N/A'}\n")
    meta.add_run(f"Overall Confidence: {document.overall_confidence or 0:.1%}\n")
    if document.hebrew_percentage:
        meta.add_run(f"Hebrew Content: {document.hebrew_percentage:.1%}\n")

    doc.add_paragraph()  # Empty line

    # Process text with corrections applied
    # Use final_text if available, otherwise ocr_text
    text = document.final_text or document.ocr_text or ""

    # Parse text into paragraphs and apply corrections
    paragraphs = text.split('\n\n')

    for para_text in paragraphs:
        if not para_text.strip():
            continue

        # Create paragraph with right alignment
        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para.paragraph_format.direction = 'RTL' if hasattr(para.paragraph_format, 'direction') else None

        # Split into words and check for corrections
        words = para_text.split()
        position = 0  # Track word position

        for word in words:
            # Check if this word has a correction
            corrected_word = corrections_map.get(position, word)

            # Add word to paragraph
            run = para.add_run(corrected_word + " ")

            # Highlight corrected words (yellow background)
            if position in corrections_map:
                try:
                    from docx.shared import RGBColor
                    run.font.highlight_color = 7  # 7 = YELLOW in Word
                except:
                    pass  # Highlight not available in this version

            position += 1

    # Add corrections summary at the end
    if errors:
        doc.add_page_break()
        summary = doc.add_paragraph()
        summary.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        summary_run = summary.add_run("Corrections Summary")
        summary_run.bold = True
        summary_run.font.size = Pt(14)

        corrected_count = sum(1 for e in errors if e.status == "corrected")
        approved_count = sum(1 for e in errors if e.status == "approved")

        stats = doc.add_paragraph()
        stats.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        stats.add_run(f"Total items reviewed: {len(errors)}\n")
        stats.add_run(f"Corrected: {corrected_count}\n")
        stats.add_run(f"Approved as-is: {approved_count}\n")

    # Save to memory buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Generate filename
    safe_filename = document.filename.replace('.pdf', '').replace('.PDF', '')
    filename = f"{safe_filename}_corrected.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{filename}"
        }
    )
